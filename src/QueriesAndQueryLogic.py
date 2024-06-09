import sys
import platform

# Monkey-patch pwd module on Windows to avoid "ImportError: No module named pwd"
# for langchain.document_loaders
if platform.system() == 'Windows':
    import pwd_compat
    sys.modules['pwd'] = pwd_compat.pwd

import PyPDF2
from langchain.chains.question_answering import load_qa_chain
from langchain.llms import OpenAI
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.document_loaders import PyPDFLoader

from langchain.chains.summarize import load_summarize_chain
from langchain.chat_models import ChatOpenAI
from langchain.prompts import ChatPromptTemplate
from langchain.chains import LLMChain

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm
from datetime import datetime
import os
from dotenv import load_dotenv

import fitz
import io
from PIL import Image

load_dotenv(override=True)
os.environ["OPENAI_API_KEY"] = os.environ.get("OPEN_AI_API_KEY")

queries = [
    "What is the publication about? Is it a new method, framework or algorithm, is it an extension of a method, framework or algorithm? Is it a Literature Review? If it is none of the above reason what the publication is about.",
    "How was the evaluation conducted? Which datasets have been used? Which key performance indicators have been used?",
    "Is an implementation available? Is a link to the implementation provided? Does the publication include a git link or links to zip files? Does the publication include links to the datasets that have been used?",
    "Are the results reproducible? Does the publication include pseudocode or formalisms to describe their approach?",
    "How many grammatical errors are in the text? How many spelling errors are in the text?"]


def generate_summary(documents):
    """
  Runs a map reduce chain over documents and returns summary
  """
    load_dotenv(override=True)
    os.environ["OPENAI_API_KEY"] = os.environ.get("OPEN_AI_API_KEY")

    llm = ChatOpenAI(model_name="gpt-4", temperature=0)
    #llm = OpenAI()
    chain = load_summarize_chain(llm, chain_type="map_reduce", verbose=False)
    summary = chain.run(documents)
    return summary


def generate_short_summary(text):
    load_dotenv(override=True)
    os.environ["OPENAI_API_KEY"] = os.environ.get("OPEN_AI_API_KEY")

    prompt = ChatPromptTemplate.from_template("Summarize this text in no more than 25 words: {text}?")
    llm = ChatOpenAI(model_name="gpt-4", temperature=0)
    #llm = OpenAI()
    chain = LLMChain(llm=llm, prompt=prompt)
    response = chain.run(text)
    return response


def database_from_documents(documents):
    embeddings = OpenAIEmbeddings()
    db = FAISS.from_documents(documents, embeddings)
    return db


def get_metadata_from_llm(db, keys):
    load_dotenv(override=True)
    os.environ["OPENAI_API_KEY"] = os.environ.get("OPEN_AI_API_KEY")

    llm_metadata = {}
    llm = ChatOpenAI(model_name="gpt-4", temperature=0)
    #llm = OpenAI()
    chain = load_qa_chain(llm=llm, chain_type="stuff")

    for key in keys:
        query = f"Give me the {key} of this paper. Do not reply in a sentence, reply only with the information. If you don't know, reply with 'not found'."
        relevant_docs = db.similarity_search(key, k=4)
        meta = chain.run(input_documents=relevant_docs, question=query)
        llm_metadata[key] = meta
    return llm_metadata


def answer_query(documents, query):
    load_dotenv(override=True)
    os.environ["OPENAI_API_KEY"] = os.environ.get("OPEN_AI_API_KEY")

    llm = ChatOpenAI(model_name="gpt-4", temperature=0)
    #llm = OpenAI()
    chain = load_qa_chain(llm=llm, chain_type="stuff")
    answer = chain.run(input_documents=documents, question=query)
    return answer


def answer_queries(db, queries):
    """
    Takes a list of documents and a list of queries.
    Gets embeddings for documents, then for each query, relevant documents are retrieved.
    Relevant documents are then sent to llm for question answering.
    Returns a list of tuples with query, answer and page numbers where relevant documents were found
    """

    queries_and_answers = []
    for query in queries:
        relevant_docs = db.similarity_search(query, k=4)

        # collect page numbers where relevant chunks are found
        found_on_pages = set()
        for doc in relevant_docs:
            found_on_pages.add(int(doc.metadata["page"]) + 1)

        # run chain
        answer = answer_query(relevant_docs, query)

        # append a tuple for each query to the list
        queries_and_answers.append((query, answer, found_on_pages))

    return queries_and_answers


def save_text_in_doc(reviews: list, output_directory, updateGUI_func, multi_output=False):
    """
    Takes a list of review dictionaries, formats them for readability and saves to docx document.
    If multi_output is set to true, one docx is generated for each review.
    If set to false, all reviews are saved in one docx document.
    """
    updateGUI_func("Saving review(s) in docx...")
    print("Saving review(s) in docx...")
    # create docx document
    if not multi_output:
        doc = Document()
    for review in reviews:
        if multi_output:
            doc = Document()
        # Add file name to document
        p = doc.add_paragraph()
        p.add_run("File: ").bold = True
        p.add_run(review["header"]["file_name"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph()
        p.add_run("Title: ").bold = True
        p.add_run(review["header"]["title"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph()
        p.add_run("Authors: ").bold = True
        p.add_run(review["header"]["authors"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph()
        p.add_run("Publication Date: ").bold = True
        p.add_run(review["header"]["publication date"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph()
        p.add_run("Keywords: ").bold = True
        p.add_run(review["header"]["keywords"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph()
        p.add_run("Words: ").bold = True
        p.add_run(str(review["header"]["number_of_words"]))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph()
        p.add_run("Pages: ").bold = True
        p.add_run(str(review["header"]["number_of_pages"]))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph()
        p.add_run("Estimated read time (minutes): ").bold = True
        p.add_run(str(review["header"]["read_time"]))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Add an additional line break
        doc.add_paragraph()

        # Add short summary to document
        p = doc.add_paragraph()
        p.add_run("Short Summary:").bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph(review["body"]["short_summary"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Add summary to document
        p = doc.add_paragraph()
        p.add_run("Summary:").bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        p = doc.add_paragraph(review["body"]["summary"])
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Add an additional line break
        doc.add_paragraph()

        # Add queries and answers to the document
        for query, answer, found_on_page in review["body"]["queries_and_answers"]:
            # Add query
            p = doc.add_paragraph()
            p.add_run(query).bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            # Add answer
            p = doc.add_paragraph(answer)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add answer
            p = doc.add_paragraph(f"Found on page(s): {found_on_page}")
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

            # Add an additional line break after an answer
            doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("Images:").bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        for image in review["images"]:
            try:
                doc.add_picture(image, width=Cm(14))
            except Exception as e:
                p = doc.add_paragraph("Couldn't save extracted image")
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p = doc.add_paragraph(str(e))
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                print("Couldn't save extracted image")
                print(e)

        if not multi_output:
            doc.add_page_break()

        timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')

        if multi_output:
            # Save the document
            doc.save(f'{output_directory}/{review["header"]["file_name"]}_review_{timestamp}.docx')

    if not multi_output:
        timestamp = datetime.now().strftime('%Y-%m-%d_%H-%M-%S')
        doc.save(f"{output_directory}/reviews{timestamp}.docx")
        print(f"file saved to {output_directory}/reviews{timestamp}.docx")
        updateGUI_func(f"file saved to {output_directory}/reviews{timestamp}.docx")

    print("Done!")


def get_documents_from_pdf(pdf_path, chunk_size_chars, overlap):
    """
    returns split documents from pdf with adjustable chunk size and overlap.
    useful for getting apropriatetly sized chunks for different tasks like summarization (larger) and question answering (smaller)
    """
    loader = PyPDFLoader(pdf_path)
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size_chars,
        chunk_overlap=overlap
    )
    documents = loader.load_and_split(text_splitter=splitter)
    return documents


def review_papers_and_save(paths, queries, output_directory, updateGUI_func):
    reviews = review_papers(paths, queries, updateGUI_func)
    save_text_in_doc(reviews, output_directory, updateGUI_func)


def review_papers(paths, queries, updateGUI_func):
    """
    Takes a list of PDF paths, reviews all of them and returns a list of review dictionaries.
    """
    updateGUI_func("Review started")
    print("Review started")

    reviews = []
    iteration = 0
    for path in paths:
        iteration += 1
        updateGUI_func(f"Review {iteration} of {len(paths)} in progress...")
        print(f"Review {iteration} of {len(paths)} in progress...")
        try:
            reviews.append(review_paper(path, queries, updateGUI_func))
            updateGUI_func(f"Review {iteration} of {len(paths)} done.")
            print(f"Review {iteration} of {len(paths)} in progress...")
        except Exception as e:
            updateGUI_func(f"Error processing paper:{path} " + str(e))
            print(f"Error processing paper:{path} ")
            print(e)

            pass

    return reviews


def review_paper(path, queries, updateGUI_func):
    """
    summarizes paper and answers questions, gathers useful metadata.
    outputs review as structured information in dict that can be used by functions to save to different output formats.
    """
    # get file name
    paper_file_name = os.path.basename(path)
    updateGUI_func(f"Reviewing file: {paper_file_name}")
    print(f"Reviewing file: {paper_file_name}")

    # count pages and words, etimate reading time
    words, pages = count_words_and_pages(path)
    read_time = round(words / 250, 1)

    # create vector database for retrieval
    updateGUI_func("Creating vector database...")
    print("Creating vector database...")
    documents = get_documents_from_pdf(path, 1000, 200)
    print(documents)
    db = database_from_documents(documents)

    # extract existing metadata from pdf
    updateGUI_func("Extracting metadata from pdf...")
    print("Extracting metadata from pdf...")
    keys_to_extract = ["title", "publication date", "keywords", "authors"]
    pdf_metadata = read_pdf_metadata(path)
    matching_pdf_meta, missing_keys = search_for_keys(keys_to_extract, pdf_metadata)

    # ask llm about missing metadata
    if (missing_keys):
        updateGUI_func("Trying to get missing metadata from llm...")
        print("Trying to get missing metadata from llm...")
        #ask specifically about title and only send the first two document, if title was not in metadata. This is to avoid retrieval of the wrong documents, assuming that the title is in the first two documents.
        if "title" in missing_keys:
            title= answer_query([documents[0],documents[1]],"what is the title of this text? reply only with the title, not in a sentence. If you don't know, reply with 'not found'.")
            #remove title from missing keys
            missing_keys.remove("title")
            #add title to found metadata
            matching_pdf_meta.update({"title":title})
        # ask about remaining missing metadata with retrieval
        llm_metadata = get_metadata_from_llm(db, missing_keys)

    # generate long summary
    updateGUI_func("Generating summary. This can take a while...")
    print("Generating summary. This can take a while...")
    summary = generate_summary(get_documents_from_pdf(path, 4000, 200))

    # generate short summary from long summary
    updateGUI_func("Generating short summary...")
    print("Generating short summary...")
    short_summary = generate_short_summary(summary)

    # answer user quries
    updateGUI_func("Answering queries. this can take a while...")
    print("Answering queries. this can take a while...")
    queries_and_answers = answer_queries(db, queries)

    images = []
    try:
        updateGUI_func("Extracting images...")
        print("Extracting images...")
        images = extract_images(path)
    except Exception as e:
        updateGUI_func("Image extraction failed: " + str(e))
        print("Image extraction failed: ")
        print(e)
        pass

    review = {
        "header": {
            "file_name": paper_file_name,
            "title": "",
            "authors": "",
            "publication date": "",
            "keywords": "",
            "number_of_words": words,
            "number_of_pages": pages,
            "read_time": read_time,
        },
        "body": {
            "summary": summary,
            "short_summary": short_summary,
            "queries_and_answers": queries_and_answers,
        },
        "images": images,
    }
    # merge metadata from pdf and llm into header
    if (matching_pdf_meta):
        review["header"].update(matching_pdf_meta)
    if (llm_metadata):
        review["header"].update(llm_metadata)

    return review


def insert_suffix(path, suffix):
    path = list(os.path.splitext(path))
    path.insert(-1, suffix)
    path = "".join(path)
    return (path)


def count_words_and_pages(pdf_path):
    # Open the PDF file in binary read mode
    pdf_file = open(pdf_path, 'rb')

    # Create a PDF reader object
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    # Initialize an empty string to store the extracted text
    raw_text = ""

    # Loop through each page and extract the text
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        raw_text += page.extract_text()
    page_count = len(pdf_reader.pages)
    # Close the PDF file
    pdf_file.close()
    word_count = len(raw_text.split())
    # return raw text
    return word_count, page_count


# currently not in use (I think):

def load_raw_text_from_pdf(path):
    """
    Reads all text from PDF and returns it as string
    """
    # Open the PDF file in binary read mode
    pdf_file = open(path, 'rb')

    # Create a PDF reader object
    pdf_reader = PyPDF2.PdfReader(pdf_file)

    # Initialize an empty string to store the extracted text
    raw_text = ""

    # Loop through each page and extract the text
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        raw_text += page.extract_text()

    # Close the PDF file
    pdf_file.close()

    # return raw text
    return raw_text


def split_raw_text(raw_text):
    """
    splits raw text and returns list of documents
    """
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    texts = text_splitter.split_text(raw_text)
    documents = text_splitter.create_documents(texts)
    # print(f"{len(documents)} documents.")
    return documents


def read_pdf_metadata(pdf_path):
    """
    Reads PDF metadata and returns it as a dictionary.
    """
    with open(pdf_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)

        # Get document information
        document_info = pdf_reader.metadata

    return document_info


def search_for_keys(keys, document_info):
    """
    Matches a list of keys against a dictionary.
    If a key is found in the dictionary and the value is not empty, that key and value is added to the output dictionary.
    If the key is not found or does not contain a value, the key is added to a set of missing keys.
    Returns a tuple containing the output dictionary and a set of missing keys.
    """
    found_meta = {}
    missing_keys = set()
    # Iterating through each string in keys
    for key_to_search in keys:
        # Check if any key in document_info contains the string
        for key in document_info:
            if key_to_search.lower() in key.lower() and document_info[key]:
                found_meta[key_to_search] = document_info[key]

            else:
                # add missing key to set
                missing_keys.add(key_to_search)

    return found_meta, missing_keys


def extract_images(path):
    # Open the PDF file
    pdf_document = fitz.open(path)
    images = []
    # Iterate through each page in the PDF
    for page_num in range(pdf_document.page_count):
        page = pdf_document[page_num]
        image_list = page.get_images(full=True)

        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_document.extract_image(xref)
            image_data = base_image["image"]

            # append file-like objects to list of images
            image = io.BytesIO(image_data)
            images.append(image)
    # Close the PDF document
    pdf_document.close()
    return images
